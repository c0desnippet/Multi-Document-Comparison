# # require fine-tuning example for dynamic 2-5 (Comparison_1_2024-10-02_12-11-53 (duplicated instead of combining))
# import os
# import re
# import pandas as pd
# from docx import Document
# import google.generativeai as genai
# from docx.shared import Inches
# from datetime import datetime


# # Connect to Gemini API
# GEMINI_API_KEY = 'AIzaSyCDap8yXYX8AaYZuRKwTkgV0Y0urHOXfpg'
# genai.configure(api_key=GEMINI_API_KEY)
# model = genai.GenerativeModel('gemini-1.5-flash')

# # Function to read a word document (ie the processed_data file)
# def read_document(file_path):
#     """
#     Read text from a DOCX document.
#     """
#     if not os.path.exists(file_path):
#         raise FileNotFoundError(f"File not found: {file_path}")
    
#     doc = Document(file_path)
#     full_text = []
#     for para in doc.paragraphs:
#         full_text.append(para.text)
#     return "\n".join(full_text)

# # Function to get the patient id of the file
# def extract_patient_id(file_path):
#     """
#     Extract PatientID from the file path.
#     """
#     match = re.search(r'PatientID_(\d+)', file_path)
#     if match:
#         return match.group(1)
#     else:
#         raise ValueError("PatientID not found in file path")
    
# def extract_timestamp_from_filename(file_path):
#     # Regex to match the date format in the filename
#     match = re.search(r'_(\d{1,2}-\d{1,2}-\d{4} \d{1,2}-\d{2})', file_path)
    
#     if match:
#         date_str = match.group(1)  # Extract the matched date string
        
#         # Convert the extracted string into a datetime object
#         try:
#             timestamp = datetime.strptime(date_str, '%d-%m-%Y %H-%M')
#             return timestamp
#         except ValueError:
#             raise ValueError(f"Failed to parse timestamp from date string: {date_str}")
#     else:
#         return None

# # Function to generate changes between two documents for PAIRWISE comparison
# def generate_comparison_report(report1_text, report2_text, date1, date2):
#     """
#     Generate comparison report between two documents.
#     The comparison will use date1 and date2 to make the output clearer.
#     """
#     # Convert datetime objects to string format
#     date1_str = date1.strftime("%Y-%m-%d %H:%M")  # Format to include date and time
#     date2_str = date2.strftime("%Y-%m-%d %H:%M")  # Format to include date and time

#     prompt = (
#         f"You will be comparing multiple radiology reports from different dates by adopting pairwise comparison between 2 documents each time."
        
#         "I will provide you with steps to do so.\n\n"
#         "Step 1: Ensure that each radiology report is organized into 3 sections: Diseases Mentioned, Organs Mentioned, and Symptoms/Phenomena of Concern.\n\n"
#         "Step 2: For each section (Diseases Mentioned, Organs Mentioned, Symptoms/Phenomena of Concern), compare the report from " + date1_str + " with the one from " + date2_str + "." 
#         " Use the report from " + date1_str + " as the base of comparison. Identify changes and categorize them into the following:\n"
#         "- New developments found in the " + date2_str + " report but not in the " + date1_str + " report (e.g., new diagnoses, progression of existing conditions).\n"
#         "- Differences (e.g., technical variations, reporting style, measurement differences, or clerical errors).\n\n"
#         "Step 3: For each section, present the findings in a table with the following columns:\n"
#         "- Category (New Development or Difference)\n"
#         "-" + date1_str + " Report Content\n"
#         "-" + date2_str + " Report Content\n"
#         "- Explanation\n\n"
#         "Report from " + date1_str + ":\n" + report1_text + "\n\n"
#         "Report from " + date2_str + ":\n" + report2_text
#     )

#     try:
#         response = model.generate_content(prompt)
#         if hasattr(response, 'text') and response.text:
#             return response.text.strip()
#         else:
#             return "Comparison report could not be generated."
#     except Exception as e:
#         print(f"Error generating comparison report: {e}")
#         return "Error generating comparison report."


# # New function to obtained FULL comparison result comprising of PAIRWISE COMPARISONS:
# def compare_multiple_reports(report_paths):
#     # Read and extract date from all reports
#     reports = [(path, read_document(path), extract_timestamp_from_filename(path)) for path in report_paths]
    
#     comparisons = []  # Store comparison results as a list of dictionaries
    
#     # Perform pairwise comparison (eg report 1 and 2, report 2 and 3, report 3 and 4, etc etc)
#     for i in range(1, len(reports)):
#         report1_path, report1_text, date1 = reports[i-1]
#         report2_path, report2_text, date2 = reports[i]
        
#         # Use dates for meaningful comparison names
#         comparison_result = generate_comparison_report(report1_text, report2_text, date1, date2)
        
#         comparisons.append({
#             'report1_path': report1_path,
#             'report2_path': report2_path,
#             'date1': date1,
#             'date2': date2,
#             'comparison_result': comparison_result
#         })
#     print(f"Comparison: {comparisons}")
#     return comparisons

# # Function to "extract" comparison results for output to word doc tables
# def parse_comparison_result(comparison_results):
#     # Initialize the structure for parsed comparison data
#     sections = {
#         # 'Diseases Mentioned': 'Diseases Mentioned',
#         # 'Organs Mentioned': 'Organs Mentioned',
#         # 'Symptoms/Phenomena of Concern': 'Symptoms/Phenomena of Concern'

#         'diseases': 'Diseases Mentioned',
#         'organs': 'Organs Mentioned',
#         'symptoms': 'Symptoms/Phenomena of Concern'
#     }

#     # Initialize a list to hold all parsed comparison data
#     all_comparison_data = []

#     # For each dictionary of pairwise comparison in the list
#     for comparison_result in comparison_results:
#         # Extract the comparison_result string from the dictionary (there are other key-item in the dictionary too, like date or path info for eg.)
#         comparison_text = comparison_result['comparison_result']  # Adjust this based on your actual key
        
#         # Extract the dates key from this pairwise comparison dictionary
#         date1 = comparison_result['date1']
#         date2 = comparison_result['date2']

#         # Initialize a dictionary to hold the parsed comparison data for the current pairwise comparison result
#         comparison_data = {key: [] for key in sections.keys()}
        
#         # Store the dates in the comparison data
#         comparison_data['date1'] = date1.strftime("%Y-%m-%d %H:%M")
#         comparison_data['date2'] = date2.strftime("%Y-%m-%d %H:%M")

#         # Adjust the section detection based on the format of the result
#         for section_key, section_title in sections.items():
#             section_start = f"{section_title}"

#             # Adjusted to capture the next section or end of text
#             section_content = re.search(f"{section_start}(.*?)(?=### |Overall|$)", comparison_text, re.DOTALL)

#             if section_content:
#                 section_text = section_content.group(1).strip()
#                 print(f"Section Text for {section_title}:\n{section_text}\n")  # Debug: Print the section content
                
#                 # Regex pattern to extract table rows (Category, Old, New, Explanation)
#                 row_pattern = re.compile(r'\|\s*(New Development|Difference)\s*\|\s*(.*?)\s*\|\s*(.*?)\s*\|\s*(.*?)\s*\|')

#                 # Extract rows in this section
#                 for match in row_pattern.finditer(section_text):
#                     category, old_content, new_content, explanation = match.groups()
#                     comparison_data[section_key].append((category.strip(), old_content.strip(), new_content.strip(), explanation.strip()))

#         # Add the parsed data for the current result, including the dates, to the overall list
#         all_comparison_data.append(comparison_data)

#     # Debug: Print all parsed comparison data. It is still a list containing dictionaries for each pairwise comparison (keys are date, sections. items are the date value, and rows of comparison within each section.)
#     print(f"Parsed Comparison Data:\n{all_comparison_data}\n")
#     return all_comparison_data


# # Function to write the comparisons to 3 main table, one for each section
# def save_multiple_comparisons(patient_id, report_paths, report_dates, comparison_result, output_folder):
#     """
#     Save the comparison results into a Word file with dynamically generated tables.
#     """
#     # Create new sub-folder if the output_folder does not have a sub-folder for the patient
#     patient_folder = os.path.join(output_folder, patient_id)
#     os.makedirs(patient_folder, exist_ok=True)

#     # Define the file_name for the compared file
#     file_name = f'Comparison_{patient_id}_{pd.Timestamp.now().strftime("%Y-%m-%d_%H-%M-%S")}.docx'
#     file_path = os.path.join(patient_folder, file_name)

#     doc = Document()
#     doc.add_heading(f'Comparison Report for Patient ID: {patient_id}', level=1)

#     # Add document paths to the report
#     doc.add_heading('Document Paths', level=2)
#     for i, path in enumerate(report_paths):
#         doc.add_paragraph(f'Report {i + 1}: {path}')

#     # Add comparison results heading
#     doc.add_heading('Comparison Results', level=2)

#     # Parse the comparison result to get data for each section
#     parsed_data = parse_comparison_result(comparison_result)
#     print(f"parsed_data: {parsed_data}")

#     # Create a dictionary to hold tables for each section
#     section_tables = {}

#     # For each section, create a table with headers for the columns
#     # sections = ['Diseases Mentioned', 'Organs Mentioned', 'Symptoms/Phenomena of Concern']
#     sections = ['diseases', 'organs', 'symptoms']

#     for section in sections:
#         doc.add_heading(f'Section: {section.capitalize()}', level=3)

#         # Create a table dynamically with columns for each report plus 2 (Category, Explanation)
#         num_reports = len(report_paths)
#         table = doc.add_table(rows=1, cols=num_reports + 2)
#         table.autofit = True
#         table.style = 'Table Grid'

#         # Add dynamic table headers
#         hdr_cells = table.rows[0].cells
#         hdr_cells[0].text = 'Category'
#         for i in range(num_reports):
#             hdr_cells[i + 1].text = f'Report {i+1} obtained on {report_dates[i].strftime("%Y-%m-%d %H:%M")}'
#         hdr_cells[num_reports + 1].text = 'Explanation'

#         # Store the table for this section
#         section_tables[section] = table

#     # parsed_data is a list containing a dictionary for each pairwise comparison
#     for comparison in parsed_data:

#         date1 = comparison['date1']
#         date2 = comparison['date2']

#         for section, rows in comparison.items():

#             if section == 'date1' or section == 'date2':
#                 continue

#             else:
#                 table = section_tables[section]  # Retrieve the existing table for the section

#                 # for each row of comparison within the section.
#                 for row_data in rows:
#                     row_cells = table.add_row().cells
#                     row_cells[0].text = row_data[0]  # Category

#                     # Dynamically populate each report's content into its respective column
#                     # from i= 0 to 3
#                     for i in range(num_reports):
#                         # Get the date for the current report from report_dates

#                         if date1 in hdr_cells[i].text:
#                             start_index = i


#                     row_cells[start_index].text = row_data[1]  # Insert Report 1 content
#                     row_cells[start_index + 1].text = row_data[2]  # Insert Report 2 content

#                     # else:
#                     #     row_cells[i + 1].text = "N/A"  # If no data is available for this report
                        
#                     row_cells[num_reports + 1].text = row_data[-1]  # Explanation

#     doc.save(file_path)

# # Updated main function to handle multiple reports
# def main():
#     report_paths = [
#         r"C:/Users/User/OneDrive - National University of Singapore/Desktop/NUS/upip/Synapxe/multi-doc/gemini_llm/pre_processing/Processed Data_attempt2/1/PatientID_1_08-7-2015 11-14.docx",  # Add up to 5 paths here
#         r"C:/Users/User/OneDrive - National University of Singapore/Desktop/NUS/upip/Synapxe/multi-doc/gemini_llm/pre_processing/Processed Data_attempt2/1/PatientID_1_09-2-2015 15-50.docx",
#         r"C:/Users/User/OneDrive - National University of Singapore/Desktop/NUS/upip/Synapxe/multi-doc/gemini_llm/pre_processing/Processed Data_attempt2/1/PatientID_1_10-2-2015 9-52.docx",
#         # r"C:/Users/User/OneDrive - National University of Singapore/Desktop/NUS/upip/Synapxe/multi-doc/gemini_llm/pre_processing/Processed Data_attempt2/1/PatientID_1_12-2-2015 19-12.docx"
        
#     ]
    
#     # Create an empty list to store report dates
#     report_dates = []
    
#     # Loop through each report path and extract the timestamp
#     for path in report_paths:
#         # Extract the date using the function
#         timestamp = extract_timestamp_from_filename(path)
        
#         if timestamp:
#             report_dates.append(timestamp)
#         else:
#             print(f"Warning: No timestamp found for {os.path.basename(path)}")

#     # Sort the reports by date and perform pairwise comparison
#     comparison_results = compare_multiple_reports(report_paths)
    
#     # Define output folder
#     output_folder = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Compared')
    
#     # Extract PatientID from first report path
#     patient_id = extract_patient_id(report_paths[0])
    
#     # Save the comparison results in a document
#     save_multiple_comparisons(patient_id, report_paths, report_dates, comparison_results, output_folder)


# if __name__ == "__main__":
#     main()




import os
import re
import pandas as pd
from docx import Document
import google.generativeai as genai
from docx.shared import Inches
from datetime import datetime


# Connect to Gemini API
GEMINI_API_KEY = 'AIzaSyCDap8yXYX8AaYZuRKwTkgV0Y0urHOXfpg'
genai.configure(api_key=GEMINI_API_KEY)
model = genai.GenerativeModel('gemini-1.5-flash')

# Function to read a word document (ie the processed_data file)
def read_document(file_path):
    """
    Read text from a DOCX document.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return "\n".join(full_text)

# Function to get the patient id of the file
def extract_patient_id(file_path):
    """
    Extract PatientID from the file path.
    """
    match = re.search(r'PatientID_(\d+)', file_path)
    if match:
        return match.group(1)
    else:
        raise ValueError("PatientID not found in file path")
    
def extract_timestamp_from_filename(file_path):
    # Regex to match the date format in the filename
    match = re.search(r'_(\d{1,2}-\d{1,2}-\d{4} \d{1,2}-\d{2})', file_path)
    
    if match:
        date_str = match.group(1)  # Extract the matched date string
        
        # Convert the extracted string into a datetime object
        try:
            timestamp = datetime.strptime(date_str, '%d-%m-%Y %H-%M')
            return timestamp
        except ValueError:
            raise ValueError(f"Failed to parse timestamp from date string: {date_str}")
    else:
        return None

# Function to generate changes between two documents for PAIRWISE comparison
def generate_comparison_report(report1_text, report2_text, date1, date2):
    """
    Generate comparison report between two documents.
    The comparison will use date1 and date2 to make the output clearer.
    """
    # Convert datetime objects to string format
    date1_str = date1.strftime("%Y-%m-%d %H:%M")  # Format to include date and time
    date2_str = date2.strftime("%Y-%m-%d %H:%M")  # Format to include date and time

    prompt = (
        f"You will be comparing multiple radiology reports from different dates by adopting pairwise comparison between 2 documents each time."
        
        "I will provide you with steps to do so.\n\n"
        "Step 1: Ensure that each radiology report is organized into 3 sections: Diseases Mentioned, Organs Mentioned, and Symptoms/Phenomena of Concern.\n\n"
        "Step 2: For each section (Diseases Mentioned, Organs Mentioned, Symptoms/Phenomena of Concern), compare the report from " + date1_str + " with the one from " + date2_str + "." 
        " Use the report from " + date1_str + " as the base of comparison. Identify changes and categorize them into the following:\n"
        "- New developments found in the " + date2_str + " report but not in the " + date1_str + " report (e.g., new diagnoses, progression of existing conditions).\n"
        "- Differences (e.g., technical variations, reporting style, measurement differences, or clerical errors).\n\n"
        "Step 3: For each section, present the findings in a table with the following columns:\n"
        "- Category (New Development or Difference)\n"
        "-" + date1_str + " Report Content\n"
        "-" + date2_str + " Report Content\n"
        "- Explanation\n\n"
        "Report from " + date1_str + ":\n" + report1_text + "\n\n"
        "Report from " + date2_str + ":\n" + report2_text
    )

    try:
        response = model.generate_content(prompt)
        if hasattr(response, 'text') and response.text:
            return response.text.strip()
        else:
            return "Comparison report could not be generated."
    except Exception as e:
        print(f"Error generating comparison report: {e}")
        return "Error generating comparison report."


# function to obtained FULL comparison result comprising of PAIRWISE COMPARISONS:
def compare_multiple_reports(report_paths):
    # Read and extract date from all reports
    reports = [(path, read_document(path), extract_timestamp_from_filename(path)) for path in report_paths]
    
    comparisons = []  # Store comparison results as a list of dictionaries
    
    # Perform pairwise comparison (eg report 1 and 2, report 2 and 3, report 3 and 4, etc etc)
    for i in range(1, len(reports)):
        report1_path, report1_text, date1 = reports[i-1]
        report2_path, report2_text, date2 = reports[i]
        
        # Use dates for meaningful comparison names
        comparison_result = generate_comparison_report(report1_text, report2_text, date1, date2)
        
        comparisons.append({
            'report1_path': report1_path,
            'report2_path': report2_path,
            'date1': date1,
            'date2': date2,
            'comparison_result': comparison_result
        })
    
    return comparisons


# Function to "extract" comparison results for output to word doc tables
def parse_comparison_result(comparison_results):
    # Initialize the structure for parsed comparison data
    sections = {
        # 'diseases': 'Diseases Mentioned',
        # 'organs': 'Organs Mentioned',
        # 'symptoms': 'Symptoms/Phenomena of Concern'
        'Diseases Mentioned': 'Diseases Mentioned',
        'Organs Mentioned': 'Organs Mentioned',
        'Symptoms/Phenomena of Concern': 'Symptoms/Phenomena of Concern'
    }

    # Initialize a list to hold all parsed comparison data
    all_comparison_data = []

    # For each dictionary of pairwise comparison in the list
    for comparison_result in comparison_results:
        # Extract the comparison_result string from the dictionary (there are other key-item in the dictionary too, like date or path info for eg.)
        comparison_text = comparison_result['comparison_result']  # Adjust this based on your actual key
        
        # Extract the dates key from this pairwise comparison dictionary
        date1 = comparison_result['date1']
        date2 = comparison_result['date2']

        # Initialize a dictionary to hold the parsed comparison data for the current pairwise comparison result
        comparison_data = {key: [] for key in sections.keys()}
        
        # Store the dates in the comparison data
        comparison_data['date1'] = date1.strftime("%Y-%m-%d %H:%M")
        comparison_data['date2'] = date2.strftime("%Y-%m-%d %H:%M")

        # Adjust the section detection based on the format of the result
        for section_key, section_title in sections.items():
            section_start = f"{section_title}"

            # Adjusted to capture the next section or end of text
            section_content = re.search(f"{section_start}(.*?)(?=### |Overall|$)", comparison_text, re.DOTALL)

            if section_content:
                section_text = section_content.group(1).strip()
                print(f"Section Text for {section_title}:\n{section_text}\n")  # Debug: Print the section content
                
                # Regex pattern to extract table rows (Category, Old, New, Explanation)
                row_pattern = re.compile(r'\|\s*(New Development|Difference)\s*\|\s*(.*?)\s*\|\s*(.*?)\s*\|\s*(.*?)\s*\|')

                # Extract rows in this section
                for match in row_pattern.finditer(section_text):
                    category, old_content, new_content, explanation = match.groups()
                    comparison_data[section_key].append((category.strip(), old_content.strip(), new_content.strip(), explanation.strip()))

        # Add the parsed data for the current result, including the dates, to the overall list
        all_comparison_data.append(comparison_data)

    # Debug: Print all parsed comparison data. It is still a list containing dictionaries for each pairwise comparison (keys are date, sections. items are the date value, and rows of comparison within each section.)
    print(f"Parsed Comparison Data:\n{all_comparison_data}\n")
    return all_comparison_data


# Function to write the comparisons to 3 main table, one for each section
def save_multiple_comparisons(patient_id, report_paths, report_dates, comparison_result, output_folder):
    """
    Save the comparison results into a Word file with dynamically generated tables.
    """
    # Create new sub-folder if the output_folder does not have a sub-folder for the patient
    patient_folder = os.path.join(output_folder, patient_id)
    os.makedirs(patient_folder, exist_ok=True)

    # Define the file_name for the compared file
    file_name = f'Comparison_{patient_id}_{pd.Timestamp.now().strftime("%Y-%m-%d_%H-%M-%S")}.docx'
    file_path = os.path.join(patient_folder, file_name)

    doc = Document()
    doc.add_heading(f'Comparison Report for Patient ID: {patient_id}', level=1)

    # Add document paths to the report
    doc.add_heading('Document Paths', level=2)
    for i, path in enumerate(report_paths):
        doc.add_paragraph(f'Report {i + 1}: {path}')

    # Add comparison results heading
    doc.add_heading('Comparison Results', level=2)

    # Parse the comparison result to get data for each section
    parsed_data = parse_comparison_result(comparison_result)

    # Create a dictionary to hold tables for each section
    section_tables = {}

    # For each section, create a table with headers for the columns
    sections = ['diseases', 'organs', 'symptoms']
    for section in sections:
        doc.add_heading(f'Section: {section.capitalize()}', level=3)

        # Create a table dynamically with columns for each report plus 2 (Category, Explanation)
        num_reports = len(report_paths)
        table = doc.add_table(rows=1, cols=num_reports + 2)
        table.autofit = True
        table.style = 'Table Grid'

        # Add dynamic table headers
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Category'
        for i in range(num_reports):
            hdr_cells[i + 1].text = f'Report {i+1} obtained on {report_dates[i].strftime("%Y-%m-%d %H:%M")}'
        hdr_cells[num_reports + 1].text = 'Explanation'

        # Store the table for this section
        section_tables[section] = table

    # parsed_data is a list containing a dictionary for each pairwise comparison
    for comparison in parsed_data:

        date1 = comparison['date1']
        date2 = comparison['date2']

        for section, rows in comparison.items():

            if section == 'date1' or section == 'date2':
                continue

            else:
                table = section_tables[section]  # Retrieve the existing table for the section

                # for each row of comparison within the section.
                for row_data in rows:
                    # print(f"Parsed Comparison Data:{row_data[1]}.")
                    row_cells = table.add_row().cells
                    row_cells[0].text = row_data[0]  # Category

                    # Dynamically populate each report's content into its respective column
                    # from i= 0 to 3
                    for i in range(num_reports):
                        # Get the date for the current report from report_dates
                        # report_date_str = report_dates[i].strftime("%Y-%m-%d %H:%M")

                        if date1 in hdr_cells[i].text:
                            start_index = i


                    row_cells[start_index].text = row_data[1]  # Insert Report 1 content
                    row_cells[start_index + 1].text = row_data[2]  # Insert Report 2 content

                    # else:
                    #     row_cells[i + 1].text = "N/A"  # If no data is available for this report
                        
                    row_cells[num_reports + 1].text = row_data[-1]  # Explanation

    doc.save(file_path)

# Updated main function to handle multiple reports
def main():
    report_paths = [
        r"C:/Users/User/OneDrive - National University of Singapore/Desktop/NUS/upip/Synapxe/multi-doc/gemini_llm/pre_processing/Processed Data_attempt2/1/PatientID_1_08-7-2015 11-14.docx",  # Add up to 5 paths here
        r"C:/Users/User/OneDrive - National University of Singapore/Desktop/NUS/upip/Synapxe/multi-doc/gemini_llm/pre_processing/Processed Data_attempt2/1/PatientID_1_09-2-2015 15-50.docx",
        r"C:/Users/User/OneDrive - National University of Singapore/Desktop/NUS/upip/Synapxe/multi-doc/gemini_llm/pre_processing/Processed Data_attempt2/1/PatientID_1_10-2-2015 9-52.docx",
    ]
    
    # Create an empty list to store report dates
    report_dates = []
    
    # Loop through each report path and extract the timestamp
    for path in report_paths:
        # Extract the date using the function
        timestamp = extract_timestamp_from_filename(path)
        
        if timestamp:
            report_dates.append(timestamp)
        else:
            print(f"Warning: No timestamp found for {os.path.basename(path)}")

    # Sort the reports by date and perform pairwise comparison
    comparison_results = compare_multiple_reports(report_paths)
    
    # Define output folder
    output_folder = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Compared')
    
    # Extract PatientID from first report path
    patient_id = extract_patient_id(report_paths[0])
    
    # Save the comparison results in a document
    save_multiple_comparisons(patient_id, report_paths, report_dates, comparison_results, output_folder)


if __name__ == "__main__":
    main()
