import os
import re
import pandas as pd
from docx import Document
import google.generativeai as genai
from docx.shared import Inches
from datetime import datetime

# Connect to Gemini API
GEMINI_API_KEY = 'AIzaSyCDap8yXYX8AaYZuRKwTkgV0Y0urHOXfpg'
genai.configure(api_key=GEMINI_API_KEY)
model = genai.GenerativeModel('gemini-1.5-flash')

# Function to read a word document (i.e., the processed_data file)
def read_document(file_path):
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return "\n".join(full_text)

# Function to get the patient id from the file path
def extract_patient_id(file_path):
    match = re.search(r'PatientID_(\d+)', file_path)
    if match:
        return match.group(1)
    else:
        raise ValueError("PatientID not found in file path")
    
def extract_timestamp_from_filename(file_path):
    match = re.search(r'_(\d{1,2}-\d{1,2}-\d{4} \d{1,2}-\d{2})', file_path)
    if match:
        date_str = match.group(1)
        try:
            timestamp = datetime.strptime(date_str, '%d-%m-%Y %H-%M')
            return timestamp
        except ValueError:
            raise ValueError(f"Failed to parse timestamp from date string: {date_str}")
    else:
        return None

# Function to generate changes between two documents for PAIRWISE comparison
def generate_comparison_report(report1_text, report2_text, date1, date2):
    date1_str = date1.strftime("%Y-%m-%d %H:%M")
    date2_str = date2.strftime("%Y-%m-%d %H:%M")

    prompt = (
        f"You will be comparing multiple radiology reports from different dates by adopting pairwise comparison between 2 documents each time."
        "I will provide you with steps to do so.\n\n"
        "Step 1: Ensure that each radiology report is organized into 3 sections: Diseases Mentioned, Organs Mentioned, and Symptoms/Phenomena of Concern.\n\n"
        "Step 2: For each section (Diseases Mentioned, Organs Mentioned, Symptoms/Phenomena of Concern), compare the report from " + date1_str + " with the one from " + date2_str + "."
        " Use the report from " + date1_str + " as the base of comparison. Identify changes and categorize them into the following:\n"
        "- New developments found in the " + date2_str + " report but not in the " + date1_str + " report (e.g., new diagnoses, progression of existing conditions).\n"
        "- Differences (e.g., technical variations, reporting style, measurement differences, or clerical errors).\n\n"
        "Step 3: For each section, present the findings in a table with the following columns:\n"
        "- Category (New Development or Difference)\n"
        "-" + date1_str + " Report Content\n"
        "-" + date2_str + " Report Content\n"
        "- Explanation\n\n"
        "Report from " + date1_str + ":\n" + report1_text + "\n\n"
        "Report from " + date2_str + ":\n" + report2_text
    )

    try:
        response = model.generate_content(prompt)
        if hasattr(response, 'text') and response.text:
            return response.text.strip()
        else:
            return "Comparison report could not be generated."
    except Exception as e:
        print(f"Error generating comparison report: {e}")
        return "Error generating comparison report."

# Compare multiple reports with the newest report as the base
def compare_multiple_reports(report_paths):
    reports = [(path, read_document(path), extract_timestamp_from_filename(path)) for path in report_paths]
    reports.sort(key=lambda x: x[2], reverse=True)
    base_report_path, base_report_text, base_date = reports[0]
    comparisons = []

    for i in range(1, len(reports)):
        report_path, report_text, report_date = reports[i]
        comparison_result = generate_comparison_report(report_text, base_report_text, report_date, base_date)
        comparisons.append({
            'report1_path': report_path,
            'report2_path': base_report_path,
            'date1': report_date,
            'date2': base_date,
            'comparison_result': comparison_result
        })
    
    return comparisons

# Parse comparison results to group similar entries
def parse_comparison_result(comparison_results):
    sections = {
        'Diseases Mentioned': 'Diseases Mentioned',
        'Organs Mentioned': 'Organs Mentioned',
        'Symptoms/Phenomena of Concern': 'Symptoms/Phenomena of Concern'
    }

    aggregated_data = {key: {} for key in sections.keys()}

    for comparison_result in comparison_results:
        comparison_text = comparison_result['comparison_result']
        date1 = comparison_result['date1'].strftime("%Y-%m-%d %H:%M")
        date2 = comparison_result['date2'].strftime("%Y-%m-%d %H:%M")

        for section_key, section_title in sections.items():
            section_start = f"{section_title}"
            section_content = re.search(f"{section_start}(.*?)(?=### |Overall|$)", comparison_text, re.DOTALL)

            if section_content:
                section_text = section_content.group(1).strip()
                row_pattern = re.compile(r'\|\s*(New Development|Difference)\s*\|\s*(.*?)\s*\|\s*(.*?)\s*\|\s*(.*?)\s*\|')

                for match in row_pattern.finditer(section_text):
                    category, old_content, new_content, explanation = match.groups()
                    common_key = old_content.strip() or new_content.strip()

                    if common_key not in aggregated_data[section_key]:
                        aggregated_data[section_key][common_key] = {
                            'category': category.strip(),
                            'entries': {}
                        }
                    
                    aggregated_data[section_key][common_key]['entries'][(date1, date2)] = (
                        old_content.strip(), new_content.strip(), explanation.strip()
                    )

    return aggregated_data

# Save aggregated comparisons to Word document
def save_multiple_comparisons(patient_id, report_paths, report_dates, comparison_result, output_folder):
    patient_folder = os.path.join(output_folder, patient_id)
    os.makedirs(patient_folder, exist_ok=True)

    file_name = f'Comparison_{patient_id}_{pd.Timestamp.now().strftime("%Y-%m-%d_%H-%M-%S")}.docx'
    file_path = os.path.join(patient_folder, file_name)

    doc = Document()
    doc.add_heading(f'Comparison Report for Patient ID: {patient_id}', level=1)
    doc.add_heading('Document Paths', level=2)
    for i, path in enumerate(report_paths):
        doc.add_paragraph(f'Report {i + 1}: {path}')

    doc.add_heading('Comparison Results', level=2)
    parsed_data = parse_comparison_result(comparison_result)

    sections = ['Diseases Mentioned', 'Organs Mentioned', 'Symptoms/Phenomena of Concern']
    for section in sections:
        doc.add_heading(f'Section: {section}', level=3)
        num_reports = len(report_paths)
        table = doc.add_table(rows=1, cols=num_reports + 3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Category'
        for i in range(num_reports):
            hdr_cells[i + 1].text = f'Report {i + 1}'
        hdr_cells[num_reports + 1].text = 'Explanation'

        section_data = parsed_data[section]
        for common_key, details in section_data.items():
            category = details['category']
            entries = details['entries']

            row_cells = table.add_row().cells
            row_cells[0].text = category

            for (date1, date2), (old_content, new_content, explanation) in entries.items():
                for i in range(num_reports - 1):
                    if date1 in hdr_cells[i + 1].text:
                        row_cells[i + 1].text = old_content or "N/A"
                        row_cells[i + 2].text = new_content or "N/A"
                        break
                row_cells[num_reports + 1].text = explanation

    doc.save(file_path)

# Main function
def main():
    report_paths = [
        r"path_to_your_report1.docx", 
        r"path_to_your_report2.docx",
        r"path_to_your_report3.docx"
    ]
    report_dates = [extract_timestamp_from_filename(path) for path in report_paths]
    comparison_results = compare_multiple_reports(report_paths)
    output_folder = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Compared')
    patient_id = extract_patient_id(report_paths[0])
    save_multiple_comparisons(patient_id, report_paths, report_dates, comparison_results, output_folder)

if __name__ == "__main__":
    main()
