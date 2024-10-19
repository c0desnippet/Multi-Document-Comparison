import os
import re
import pandas as pd
from docx import Document
import google.generativeai as genai
from docx.shared import Inches


# Connect to Gemini API
GEMINI_API_KEY = 'AIzaSyCDap8yXYX8AaYZuRKwTkgV0Y0urHOXfpg'
genai.configure(api_key=GEMINI_API_KEY)
model = genai.GenerativeModel('gemini-1.5-flash')

# Function to read a word document (ie the processed_data file)
def read_document(file_path):
    """
    Read text from a DOCX document.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return "\n".join(full_text)

# Function to get the patient id of the file
def extract_patient_id(file_path):
    """
    Extract PatientID from the file path.
    """
    match = re.search(r'PatientID_(\d+)', file_path)
    if match:
        return match.group(1)
    else:
        raise ValueError("PatientID not found in file path")
    

# Function to generate the changes between both documents
def generate_comparison_report(old_text, new_text):
    """
    Use Gemini LLM to generate a comparison report.
    """

    # #third comparison file output
    # prompt = (
    # "You are comparing two radiology reports that are already organized into sections: Diseases Mentioned, Organs Mentioned, and Symptoms/Phenomena of Concern.\n"
    # "For each section (Diseases Mentioned, Organs Mentioned, Symptoms/Phenomena of Concern), compare the old and new reports. Identify and distinguish between:\n"
    # "1. New developments (e.g., new diagnoses, progression of existing conditions, new symptoms).\n"
    # "2. Differences (e.g., technical variations, reporting style, measurement differences, or clerical errors).\n"
    # "Provide explanations for why you categorized each change as a 'New development' or a 'Difference'.\n\n"
    # "Old Report:\n" + old_text + "\n\n"
    # "New Report:\n" + new_text
    # )

    # Edited CoT prompt:
    prompt = (
        "You will be comparing two radiology reports provided. I will be providing you with the steps to do so.\n\n"
        
        "Step 1: Ensure that each radiology report is organized into 3 sections: Diseases Mentioned, Organs Mentioned, and Symptoms/Phenomena of Concern.\n\n"
        
        "Step 2: For each section (Diseases Mentioned, Organs Mentioned, Symptoms/Phenomena of Concern), compare the old report with the new report. Use the old report as the base of comparison. Identify changes and categorize them into the following:\n"
        "- New developments found in the new report but not in the old report (e.g., new diagnoses, progression of existing conditions).\n"
        "- Differences (e.g., technical variations, reporting style, measurement differences, or clerical errors).\n\n"
        
        "Step 3: For each section, present the findings in a table with the following columns:\n"
        "- Category (New Development/Difference)\n"
        "- Old Report Content\n"
        "- New Report Content\n"
        "- Explanation\n\n"
        
        "Step 4: After the table, provide an explanation for why each change was categorized as a 'New development' or a 'Difference'.\n\n"
        
        "Step 5: Repeat Steps 2-4 for the other two sections (Organs Mentioned and Symptoms/Phenomena of Concern).\n\n"
        
        "Old Report:\n" + old_text + "\n\n"
        "New Report:\n" + new_text
    )

    
    try:
        response = model.generate_content(prompt)
        
        # Check if the response contains the generated text
        if hasattr(response, 'text') and response.text:
            return response.text.strip()
        elif 'content' in response and response['content']:
            return response['content'].strip()
        else:
            return "Comparison report could not be generated."
    except Exception as e:
        print(f"Error generating comparison report: {e}")
        return "Error generating comparison report."


# # Function to save the comparison writeup in a word file
# def save_comparison_results(patient_id, old_report_path, new_report_path, comparison_result, output_folder):
#     """
#     Save the comparison results into a Word file.
#     """
#     # Create new sub-folder if the output_folder does not have a sub-folder for the patient
#     patient_folder = os.path.join(output_folder, patient_id)
#     os.makedirs(patient_folder, exist_ok=True)
    
#     # Define the file_name for the compared file
#     file_name = f'Comparison_{patient_id}_{pd.Timestamp.now().strftime("%Y-%m-%d_%H-%M-%S")}.docx'
#     file_path = os.path.join(patient_folder, file_name)
    
#     doc = Document()
#     doc.add_heading(f'Comparison Report for Patient ID: {patient_id}', level=1)
    
#     # Add document paths to the report
#     doc.add_heading('Document Paths', level=2)
#     doc.add_paragraph(f'Old Report: {old_report_path}')
#     doc.add_paragraph(f'New Report: {new_report_path}')
    
#     # Add comparison results
#     doc.add_heading('Comparison Results', level=2)
#     doc.add_paragraph(comparison_result)
    
#     doc.save(file_path)


def parse_comparison_result(comparison_result):
    """
    Parse the comparison_result text and extract section-based comparison data.
    Returns a dictionary where each section contains a list of tuples with (Category, Old Content, New Content, Explanation).
    """
    sections = {
        "Diseases Mentioned": "Diseases Mentioned",
        "Organs Mentioned": "Organs Mentioned",
        "Symptoms/Phenomena of Concern": "Symptoms/Phenomena of Concern"
    }
    comparison_data = {section: [] for section in sections}

    # Debug: Print the raw comparison result
    print(f"Raw comparison result:\n{comparison_result}\n")

    # Adjust the section detection based on the format of the result
    for section_key, section_title in sections.items():
        section_start = f"### Section {list(sections.keys()).index(section_key) + 1}: {section_title}"
        # Adjusted to capture next section or end of text
        section_content = re.search(f"{section_start}(.*?)(?=### Section|Overall|$)", comparison_result, re.DOTALL)
        
        if section_content:
            section_text = section_content.group(1).strip()
            print(f"Section Text for {section_title}:\n{section_text}\n")  # Debug: Print the section content
            
            # Regex pattern to extract table rows (Category, Old, New, Explanation)
            row_pattern = re.compile(r'\|\s*(New Development|Difference)\s*\|\s*(.*?)\s*\|\s*(.*?)\s*\|\s*(.*?)\s*\|')
            
            # Extract rows in this section
            for match in row_pattern.finditer(section_text):
                category, old_content, new_content, explanation = match.groups()
                comparison_data[section_key].append((category.strip(), old_content.strip(), new_content.strip(), explanation.strip()))

    # Debug: Print the parsed comparison data
    print(f"Parsed Comparison Data:\n{comparison_data}\n")
    
    return comparison_data


def save_comparison_results(patient_id, old_report_path, new_report_path, comparison_result, output_folder):
    """
    Save the comparison results into a Word file with dynamically generated tables.
    """
    # Create new sub-folder if the output_folder does not have a sub-folder for the patient
    patient_folder = os.path.join(output_folder, patient_id)
    os.makedirs(patient_folder, exist_ok=True)
    
    # Define the file_name for the compared file
    file_name = f'Comparison_{patient_id}_{pd.Timestamp.now().strftime("%Y-%m-%d_%H-%M-%S")}.docx'
    file_path = os.path.join(patient_folder, file_name)
    
    doc = Document()
    doc.add_heading(f'Comparison Report for Patient ID: {patient_id}', level=1)
    
    # Add document paths to the report
    doc.add_heading('Document Paths', level=2)
    doc.add_paragraph(f'Old Report: {old_report_path}')
    doc.add_paragraph(f'New Report: {new_report_path}')
    
    # Add comparison results heading
    doc.add_heading('Comparison Results', level=2)

    # Parse the comparison result to get data for each section
    parsed_data = parse_comparison_result(comparison_result)

    # Iterate over each section and add the data into a table
    for section, rows in parsed_data.items():
        doc.add_heading(f'Section: {section}', level=3)
        
        # Create a table with 4 columns: 'Category', 'Old Report Content', 'New Report Content', and 'Explanation'
        table = doc.add_table(rows=1, cols=4)
        table.autofit = True
        table.style = 'Table Grid'
        
        # Add table headers
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Category'
        hdr_cells[1].text = 'Old Report Content'
        hdr_cells[2].text = 'New Report Content'
        hdr_cells[3].text = 'Explanation'
        
        # Populate the table with actual data from the parsed result
        for row_data in rows:
            row_cells = table.add_row().cells
            row_cells[0].text = row_data[0]  # Category
            row_cells[1].text = row_data[1]  # Old Report Content
            row_cells[2].text = row_data[2]  # New Report Content
            row_cells[3].text = row_data[3]  # Explanation

    doc.save(file_path)


# Main application
def main():
    # To edit: Change the paths accordingly to the 2 files you want to compare
    old_report_path = r"C:/Users/User/OneDrive - National University of Singapore/Desktop/NUS/upip/Synapxe/multi-doc/gemini_llm/pre_processing/Processed Data_attempt2/1/PatientID_1_09-2-2015 15-50.docx"
    new_report_path = r"C:/Users/User/OneDrive - National University of Singapore/Desktop/NUS/upip/Synapxe/multi-doc/gemini_llm/pre_processing/Processed Data_attempt2/1/PatientID_1_10-2-2015 9-52.docx"
    
    # Path to the "comparing" folder
    comparing_folder = os.path.dirname(os.path.abspath(__file__))
    
    # Define the relative path to the "Compared" folder within the "comparing" folder
    output_folder = os.path.join(comparing_folder, 'Compared')
    
    old_text = read_document(old_report_path)
    new_text = read_document(new_report_path)
    
    # Generate comparison report using Gemini LLM
    comparison_result = generate_comparison_report(old_text, new_text)
    
    # Extract PatientID from file paths
    patient_id = extract_patient_id(old_report_path)
    
    save_comparison_results(patient_id, old_report_path, new_report_path, comparison_result, output_folder)

# Example usage
if __name__ == "__main__":
    main()