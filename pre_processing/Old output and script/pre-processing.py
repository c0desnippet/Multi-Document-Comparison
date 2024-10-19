# import pandas as pd
# from docx import Document
# from pathlib import Path
# import os

# # Load the CSV file
# df = pd.read_csv('Chest Scans_deidentified.csv')  # Replace with your file name

# # Define the base folder for processed data
# base_folder = Path('Processed Data')
# base_folder.mkdir(exist_ok=True)

# # Iterate over each row in the DataFrame
# for index, row in df.iterrows():
#     # Extract PatientID and Updated Date
#     patient_id = str(row['Masked_PatientID'])
#     updated_date = row['Performed Date Time']  # Ensure this is the correct column name
    
#     # Create a folder for the PatientID if it doesn't exist
#     patient_folder = base_folder / patient_id
#     patient_folder.mkdir(exist_ok=True)
    
#     # Create a new Document
#     doc = Document()

#     # Add a title to the document (optional)
#     doc.add_heading(f'Row {index + 1}', level=1)

#     # Add the content of the row organized by the column name
#     for col in df.columns:
#         doc.add_paragraph(f'{col}: {row[col]}')

#     # Define the filename using PatientID and Updated Date
#     # Ensure updated_date is in a valid filename format
#     safe_date = updated_date.replace('/', '-').replace(':', '-')
#     filename = patient_folder / f'PatientID_{patient_id}_{safe_date}.docx'
    
#     # Save the document
#     doc.save(filename)

# print("Documents created and organized successfully!")

import pandas as pd
from docx import Document
from pathlib import Path
import os
import spacy

# Load the SciSpacy biomedical model for entity recognition
nlp = spacy.load('en_core_sci_md')  # Use the correct model for your biomedical text

# Load the CSV file
df = pd.read_csv('Chest Scans_deidentified.csv')  # Replace with your file name

# Define the base folder for processed data
base_folder = Path('Processed Data')
base_folder.mkdir(exist_ok=True)

# Helper function to extract entities from the report text
def extract_entities(report_text):
    doc = nlp(report_text)
    diseases = set()
    organs = set()
    symptoms = set()

    # Iterate over recognized entities
    for ent in doc.ents:
        if ent.label_ in ['DISEASE', 'CONDITION']:  # Assuming disease labels
            diseases.add(ent.text)
        elif ent.label_ == 'ORGAN':  # Assuming organ labels
            organs.add(ent.text)
        elif ent.label_ == 'SYMPTOM':  # Assuming symptom labels
            symptoms.add(ent.text)

    return diseases, organs, symptoms

# Helper function to generate layman explanation
def generate_layman_explanation(report_text):
    # You can enhance this with custom logic or external medical libraries
    explanation = f"This radiology report discusses {report_text}. In simpler terms, this means..."
    return explanation

# Helper function to generate a summary based on NLP-extracted entities
def generate_summary(report_text):
    diseases, organs, symptoms = extract_entities(report_text)

    # Format the summary
    disease_summary = "Detected diseases: " + ', '.join(diseases) if diseases else "No diseases detected."
    organ_summary = "Organs mentioned: " + ', '.join(organs) if organs else "No specific organs mentioned."
    symptom_summary = "Symptoms or phenomena: " + ', '.join(symptoms) if symptoms else "No symptoms mentioned."
    
    summary = f"{disease_summary}\n{organ_summary}\n{symptom_summary}"
    return summary

# Iterate over each row in the DataFrame
for index, row in df.iterrows():
    # Extract PatientID and Updated Date
    patient_id = str(row['Masked_PatientID'])
    updated_date = row['Performed Date Time']  # Ensure this is the correct column name
    
    # Create a folder for the PatientID if it doesn't exist
    patient_folder = base_folder / patient_id
    patient_folder.mkdir(exist_ok=True)
    
    # Create a new Document
    doc = Document()

    # Add a title to the document (optional)
    doc.add_heading(f'Row {index + 1}', level=1)

    # Add the content of the row organized by the column name
    report_text = ""
    for col in df.columns:
        value = str(row[col])
        doc.add_paragraph(f'{col}: {value}')
        # Assuming one of the columns contains the radiology report text
        if col == 'Radiology Report':  # Adjust to match actual column name
            report_text = value

    # Add the layman explanation section
    layman_explanation = generate_layman_explanation(report_text)
    doc.add_heading('Layman Explanation', level=2)
    doc.add_paragraph(layman_explanation)

    # Add the summary section
    summary = generate_summary(report_text)
    doc.add_heading('Summary', level=2)
    doc.add_paragraph(summary)

    # Define the filename using PatientID and Updated Date
    safe_date = updated_date.replace('/', '-').replace(':', '-')
    filename = patient_folder / f'PatientID_{patient_id}_{safe_date}.docx'
    
    # Save the document
    doc.save(filename)

print("Documents created and organized successfully!")
