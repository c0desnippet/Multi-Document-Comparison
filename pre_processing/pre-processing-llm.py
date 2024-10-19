import pandas as pd
from docx import Document
from pathlib import Path
import google.generativeai as genai

# Connect to Gemini API
GEMINI_API_KEY = 'AIzaSyCDap8yXYX8AaYZuRKwTkgV0Y0urHOXfpg'
genai.configure(api_key=GEMINI_API_KEY)
model = genai.GenerativeModel('gemini-1.5-flash')

# Function to generate summary based on focus questions (for acceptable memory loss)
def generate_summary(extracted_text):
    prompt = (
        "The following text is extracted from a radiology report."
        "Firstly, determine and remember what type of image the text was extracted from."
        "Next, I would like you to summerise the text extracted based on the following guiding questions:\n"
        "1. Any disease(s) mentioned in the radiology report? If yes, include all elaboration related to the disease(s)."
        "Do not interpret the disease name from the text. If no disease name mentioned, just leave the section as NIL.\n"
        "2. Any organ(s) mentioned in the radiology report? If yes, include all information regarding the organ(s)."
        # "Do not interpret the organ name from the text. If no organs mentioned, just leave the section as NIL.\n"
        "3. Any symptoms or phenomenon that would cause attention? If yes, please elaborate on the concerns.\n"
        # "Do not include any extraneous information or explanations outside of the report provided.\n"
        f"Text: {extracted_text}"
    )

    try:
        response = model.generate_content(prompt)
        
        # Debugging: Print the raw response
        print("API Response:", response)
        
        # Safely access response content
        if hasattr(response, 'text') and response.text:
            return response.text.strip()
        elif 'content' in response and response['content']:
            return response['content'].strip()
        else:
            return "Summary could not be generated."
    except Exception as e:
        print(f"Error generating summary: {e}")
        return "Error generating summary."

# Function to generate laymen explantion of the report
def generate_layman_explanation(extracted_text):
    prompt = (
        "The following text is extracted from a radiology report."
        "You are an interpreter tasked to translate the radiology report 'Text' section into layman terms./n"
        "Remember that your audience do not have any prior medical knowledge.\n"
        "Refrain from using medically intensive jargons.\n"
        "Do not include any extraneous information or explanations. Provide a complete, clear, and concise layman summary of the extracted content.\n"
        f"Text: {extracted_text}"
    )

    try:
        response = model.generate_content(prompt)
        
        # Debugging: Print the raw response
        print("API Response:", response)
        
        # Safely access response content
        if hasattr(response, 'text') and response.text:
            return response.text.strip()
        elif 'content' in response and response['content']:
            return response['content'].strip()
        else:
            return "Summary could not be generated."
    except Exception as e:
        print(f"Error generating summary: {e}")
        return "Error generating summary."

# To edit accordingly: Load the raw CSV file containing rows of radiology "reports"
df = pd.read_csv('Chest Scans_deidentified.csv')  # Replace with your file name
print("read csv")

# To edit accordingly: Define the NAME of the base folder to store processed data files
base_folder = Path('Processed Data_attempt2')
base_folder.mkdir(exist_ok=True)

# Iterate over each row in the DataFrame
for index, row in df.iterrows():
    # Extract PatientID and Updated Date
    patient_id = str(row['Masked_PatientID'])
    updated_date = row['Performed Date Time']  # Ensure this is the correct column name
    
    # Create a sub-folder within the "Processed_Data_x" folder for the PatientID if it doesn't exist
    patient_folder = base_folder / patient_id
    patient_folder.mkdir(exist_ok=True)
    
    # Create a new Document
    doc = Document()

    # Add a title to the document using PatientID and Updated Date
    doc.add_heading(f'Patient ID: {patient_id}, Performed Date: {updated_date}', level=1)

    # Add the raw report section
    doc.add_heading('Raw Radiology Report Extracted', level=2)

    # Add the content of the row organized by the column name
    report_text = ""
    for col in df.columns:
        value = str(row[col])
        doc.add_paragraph(f'{col}: {value}')
        # Assuming one of the columns contains the radiology report text
        if col == 'Text':  # Adjust to match actual column name
            report_text = value

    # Add the layman explanation section
    layman_explanation = generate_layman_explanation(report_text)
    doc.add_heading('Layman Explanation', level=2)
    doc.add_paragraph(layman_explanation)

    # Add the summary section
    summary = generate_summary(report_text)
    doc.add_heading('Summary', level=2)
    doc.add_paragraph(summary)

    # Define the filename using PatientID and Updated Date
    safe_date = updated_date.replace('/', '-').replace(':', '-')
    filename = patient_folder / f'PatientID_{patient_id}_{safe_date}.docx'
    
    # Save the document
    doc.save(filename)

print("Documents created and organized successfully!")
